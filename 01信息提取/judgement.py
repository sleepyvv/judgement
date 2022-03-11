import docx
import re
import os
import xlsxwriter
from random import shuffle
#获取文档对象
# 定义标点符号和特殊字母
punctuation = '''，。、:；（）ＸX×xa"“”,<《》'''
fayuan="人民法院"
panjue="判决如下"
panjue_end="审判长"
panjue_end3="审　判　长"
panjue_end2="审判员"
panjue_end4="审　判　员"
end_line=999999
panjue_line=999999
fact="本院认为"
zuigao="最高人民法院"
fact_line=0
number=0
test="黑龙江省大庆市红岗区"
path="C:\\Users\\lenovo\\Desktop\\test"
filelist = [path+"\\" + i for i in os.listdir(path)]
excel_pos=2
data=open("C:\\Users\\lenovo\\Desktop\\test\\data.txt",'w+',encoding='utf-8') 
workbook = xlsxwriter.Workbook('C:\\Users\\lenovo\\Desktop\\test\\hello.xlsx') # 建立文件
worksheet = workbook.add_worksheet() # 建立sheet， 可以work.add_worksheet('employee')来指定sheet名，但中文名会报UnicodeDecodeErro的错误
worksheet.write('A1', 'Hello world') # 向A1写入
#worksheet.write(excel_pos,1,'最高人民法院')#向第二行第二例写入guoshun
for file in filelist:
    if file.endswith (".docx"):
        now_file=file.replace('\\','\\\\')
        fopen=docx.Document(now_file)
#        print(now_file)
#        print("法院为："+fopen.paragraphs[0].text)
        fayuan_location=fopen.paragraphs[0].text
        zuigao_pos=fayuan_location.find('最高人民法院')
        city_tmp=fayuan_location.find('市')
        province_tmp=fayuan_location.find('省')
        if province_tmp==-1:
            province_tmp=fayuan_location.find('壮族自治区')
            if province_tmp==-1:
                province_tmp=fayuan_location.find('维吾尔自治区')
        niubi=fayuan_location.find('中华人民共和国')    
    #    print(niubi)
        if zuigao_pos>0:
            print("最高人民法院",end="",file=data)
            #worksheet.write(excel_pos,1,'最高人民法院')#向第二行第二例写入guoshun
        elif province_tmp > 0:
            if(niubi==-1):
                print(fayuan_location[0:province_tmp],end="",file=data)
                worksheet.write(excel_pos,1,fayuan_location[0:province_tmp])
            else:
                print(fayuan_location[niubi+7:province_tmp],end="",file=data)
                worksheet.write(excel_pos,1,fayuan_location[niubi+7:province_tmp])
        elif city_tmp>0 : 
            if(niubi==-1):
                print(fayuan_location[0:city_tmp],end="",file=data)
                worksheet.write(excel_pos,1,fayuan_location[0:city_tmp])
            else:
                print(fayuan_location[niubi+7:city_tmp],end="",file=data)
                worksheet.write(excel_pos,1,fayuan_location[niubi+7:city_tmp])
        print(" ",end="",file=data)
        excel_pos=excel_pos+1
        for i in range(len(fopen.paragraphs)):
            if fact in fopen.paragraphs[i].text:
              panjue_line=i
            if (panjue_end in fopen.paragraphs[i].text) or (panjue_end3 in fopen.paragraphs[i].text):
               panjue_line=99999
            if panjue_end2 in fopen.paragraphs[i].text or (panjue_end4 in fopen.paragraphs[i].text):
               panjue_line=99999
            if i >= panjue_line:
                print(fopen.paragraphs[i].text,end="",file=data)
        print(" ",file=data)
        number=number+1
        print("file done ",number)
workbook.close()
"""        else:
            print(now_file)
            print("问题："+fayuan_location)
            os.remove(now_file)
      #  print("判决是：")
"""
""" 
       for i in range(len(fopen.paragraphs)):
            if panjue in fopen.paragraphs[i].text:
                panjue_line=i
            if panjue_end in fopen.paragraphs[i].text:
                panjue_line=99999
            if i > panjue_line:
                print(fopen.paragraphs[i].text)
"""
"""
file=docx.Document("C:\\Users\\lenovo\\Desktop\\1.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
#输出每一段的内容
#for para in file.paragraphs:
#    print(para.text)
print("法院为："+file.paragraphs[0].text)
fayuan_location=file.paragraphs[0].text
tmp=fayuan_location.find('市')
print("地区："+ fayuan_location[0:tmp])
#print  url[0:url.rfind('/v2')]
#输出段落编号及段落内容
print("判决是：")
for i in range(len(file.paragraphs)):
    if panjue in file.paragraphs[i].text:
        panjue_line=i
    if panjue_end in file.paragraphs[i].text:
        panjue_line=99999
    if i > panjue_line:
     print(file.paragraphs[i].text)
"""