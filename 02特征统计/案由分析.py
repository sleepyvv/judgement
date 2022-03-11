import docx
import re
import os
import xlsxwriter
path="C:\\Users\\lenovo\\Desktop\\test"
filelist = [path+"\\" + i for i in os.listdir(path)]
filenamelist=[i for i in os.listdir(path)]
print(filenamelist)
gsly_count=0
ht_count=0
gd_count=0
yspj=0
yscd=0
espj=0
escd=0
zspj=0
zscd=0
yuangao="原告"
gongsi="公司"
gsyg=0
gsbg=0
ysgs=0
esgs=0
dsdl=0
bohuigs=0
esbohuigs=0
for file in filelist:
    ifgsyg=0
    if file.find('损害公司利益') >1:
        gsly_count=gsly_count+1
    if file.find('合同纠纷') >1:
        ht_count=ht_count+1
    if file.find('股东利益')>1:
        gd_count=gd_count+1
    if file.find('一审')>=1:
        if file.find('判决')>=1:
            yspj=yspj+1
            now_file=file.replace('\\','\\\\')
            fopen=docx.Document(now_file)
            for i in range(len(fopen.paragraphs)):
                if yuangao in fopen.paragraphs[i].text:
                    if gongsi in fopen.paragraphs[i].text:
                        gsyg=gsyg+1
                        ysgs=ysgs+1
                        ifgsyg=1
                        for j in range(i,len(fopen.paragraphs)):
                            if "代表人" in fopen.paragraphs[j].text and "董事" in fopen.paragraphs[j].text:
                                dsdl=dsdl+1
                                break
                            if "代表人" in fopen.paragraphs[j].text and "股东" in fopen.paragraphs[j].text:
                                print(now_file)
                                break
                            if "代表人" in fopen.paragraphs[j].text and "监事" in fopen.paragraphs[j].text:
                                print(now_file)
                                break
                        break
                    elif "董事会" in fopen.paragraphs[i].text:
                        print("我超！董事会！")
                    elif "监事会" in fopen.paragraphs[i].text:
                        print("我超！监事会！")
                        break
                    elif "股东" in fopen.paragraphs[i].text:
                        print("我超！股东子！")
                        print(now_file)
                        break
                    else:
                        gsbg=gsbg+1
                        break
            for i in range(len(fopen.paragraphs)):
                if "驳回上诉" in fopen.paragraphs[i].text:
                    if ifgsyg==1 :
                        #print("驳回公司上诉")
                        bohuigs=bohuigs+1
                        break
        elif file.find('裁定')>=1:
            yscd=yscd+1
    if file.find('二审')>=1:
          if file.find('判决')>=1:
            espj=espj+1
            now_file=file.replace('\\','\\\\')
            fopen=docx.Document(now_file)
            for i in range(len(fopen.paragraphs)):
                if yuangao in fopen.paragraphs[i].text:
                    if gongsi in fopen.paragraphs[i].text:
                        gsyg=gsyg+1
                        esgs=esgs+1
                        ifgsyg=1
                        for p in range(i,len(fopen.paragraphs)):
                            if "代表人" in fopen.paragraphs[p].text and "董事" in fopen.paragraphs[p].text:
                                dsdl=dsdl+1
                                break
                            if "代表人" in fopen.paragraphs[p].text and "股东" in fopen.paragraphs[p].text:
                                print(now_file)
                                break
                            if "代表人" in fopen.paragraphs[p].text and "监事" in fopen.paragraphs[p].text:
                                print(now_file)
                                break
                        break
                    else:
                        gsbg=gsbg+1
                        break
            for i in range(len(fopen.paragraphs)):
                if "维持原判" in fopen.paragraphs[i].text:
                    if ifgsyg==1 :
                        #print("驳回公司上诉")
                        esbohuigs=esbohuigs+1
                        break
          elif file.find('裁定')>=1:
            escd=escd+1
    if file.find('再审')>=1:
        if file.find('判决')>=1:
            zspj=zspj+1
        elif file.find('裁定')>=1:
            zscd=zscd+1

            
print("损害公司利益有：",end="")
print(gsly_count)
print("合同纠纷有：",end="")
print(ht_count)
print("股东利益有：",end="")
print(gd_count)
print(yspj)
print(yscd)
print(espj)
print(escd)
print(zspj)
print(zscd)
print("公司原告数："+str(gsyg)+"  一审："+str(ysgs)+"  二审 ：" +str(esgs))
print("公司被告数："+str(gsbg))
print("驳回公司上诉数"+str(bohuigs))
print("二审驳回公司上诉数"+str(esbohuigs))
print(dsdl)

        